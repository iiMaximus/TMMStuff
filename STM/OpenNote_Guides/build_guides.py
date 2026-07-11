from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
TEXT = ROOT / "_source_text"
MD_OUT = ROOT / "markdown"
DOCX_OUT = ROOT / "docx"
ASSETS = ROOT / "assets"


def clean_line(line: str) -> str:
    line = line.replace("\u2013", "-").replace("\u2014", "-").replace("\u2212", "-")
    line = re.sub(r"\s+", " ", line).strip()
    return line


def source_appendix(text_file: str, max_lines_per_page: int = 18) -> str:
    path = TEXT / text_file
    if not path.exists():
        return ""
    raw = path.read_text(encoding="utf-8", errors="ignore")
    pages = re.split(r"\n--- PAGE (\d+) ---\n", raw)
    out = ["## Source Slide Checklist", "", "This is a compact OCR/text extraction from the source PDF. Use it as a searchable reminder of the slide wording, not as a replacement for diagrams.", ""]
    for i in range(1, len(pages), 2):
        page_no = pages[i]
        body = pages[i + 1]
        lines = []
        for line in body.splitlines():
            line = clean_line(line)
            if not line or line == page_no:
                continue
            if line.startswith("http"):
                continue
            if re.fullmatch(r"\d+", line):
                continue
            lines.append(line)
        if not lines:
            continue
        if len(lines) > max_lines_per_page:
            keep = lines[:max_lines_per_page]
            keep.append("...")
            lines = keep
        out.append(f"### Source page {page_no}")
        out.extend(f"- {line}" for line in lines)
        out.append("")
    return "\n".join(out).strip() + "\n"


def section(title: str, body: str) -> str:
    return f"## {title}\n\n{body.strip()}\n\n"


def guide(title: str, source: str, body: str, text_file: str | None = None, appendix_lines: int = 18) -> str:
    parts = [
        f"# {title}",
        "",
        f"Source: `{source}`",
        "",
        "Purpose: open-note exam guide. It prioritizes answer structure, formulas, tutorial methods, and professor-style traps.",
        "",
        body.strip(),
        "",
    ]
    if text_file:
        parts.append(source_appendix(text_file, appendix_lines))
    return "\n".join(parts).strip() + "\n"


GUIDES: list[dict[str, str]] = []


GUIDES.append({
    "slug": "00_master_formula_and_exam_methods",
    "title": "Master Formula and Exam Methods",
    "source": "STM/Material(1), STM/Material(2), Exam simulation_STM eng.pdf",
    "md": guide(
        "Master Formula and Exam Methods",
        "STM/Material(1), STM/Material(2), Exam simulation_STM eng.pdf",
        section("Exam Pattern To Prepare For", """
- The exam simulation is 45 minutes with 9 questions. It mixes open theory, multiple choice, and calculations.
- High-probability calculation families: tensile/elastic design, thermal expansion or thermal stress, generic phase diagrams, and Fe-C diagram calculations.
- High-probability theory families: heat conduction, sintering, ceramic porosity, glass/safety glass, thermoplastic vs thermoset polymers, mechanical test interpretation, and material-class comparisons.
- In multiple choice, the professor often asks for the only correct item. Eliminate options that use absolute wording such as "no effect", "always", or "only" when the lecture taught a conditional mechanism.
""")
        + section("Mechanical Formula Block", """
- Stress: sigma = F / A. Use N/mm2 = MPa when force is in N and area is in mm2.
- Engineering strain: e = (l - l0) / l0 = Delta l / l0. It is dimensionless; multiply by 100 for percent.
- Hooke law in the elastic field: sigma = E e. Therefore e = sigma / E and Delta l = (sigma / E) l0.
- Circular section: A = pi r2 = pi d2 / 4.
- Safety factor for no plastic deformation: sigma_allowed = YS / n, then Fmax = sigma_allowed A.
- Poisson effect: nu = - transverse strain / longitudinal strain. If a diameter decreases, transverse strain is negative and tensile longitudinal strain is positive.
- Seemly/engineering fracture stress uses the initial area. True fracture stress uses the final reduced area at rupture.
- Ductility or elongation at fracture: A% = 100 (lf - l0) / l0.
- Necking/reduction of area: Z% = 100 (A0 - Af) / A0. For a circular section, area scales with d2.
- Toughness: area under the stress-strain curve up to fracture. Resilience modulus in the elastic field: Ur = 1/2 sigma_y e_y = sigma_y2 / (2E).
- Fracture toughness: KIc = Y sigma_c sqrt(pi a). Failure occurs when K reaches KIc. Larger cracks, higher stress, and unfavorable geometry increase risk.
- Crack-tip stress concentration for an elliptical crack: sigma_m = 2 sigma_0 sqrt(a / r_t). Sharper cracks have smaller r_t and higher local stress.
- Hardness examples: Brinell HB = F / (pi D d). Vickers HV = 1.854 F / L2. For steels, a rough correlation is TS about 3.45 HB.
""")
        + section("Creep and Fatigue Block", """
- Creep test: constant load at constant high temperature; record strain versus time.
- Creep becomes important for metals above about 0.3-0.4 Tm, ceramics above about 0.6-0.7 Tm, and amorphous materials above Tg.
- Creep stages: primary has decreasing strain rate; secondary has approximately constant strain rate; tertiary accelerates because voids and cracks form until rupture.
- Higher stress and higher temperature shift creep curves upward and reduce life.
- Creep resistance improves with high Tm/high E materials, large grains or single crystals, solid-solution strengthening, precipitates, and second phases.
- Fatigue is failure under cyclic loading, often below YS or UTS. Use an S-N/Wohler curve to read fatigue life or fatigue strength.
- Stress amplitude: sigma_a = (sigma_max - sigma_min) / 2.
- Ferrous alloys can show a fatigue limit; many non-ferrous alloys such as Al and Cu do not show a true intrinsic limit in the same way.
""")
        + section("Thermal Formula Block", """
- Heat capacity: C = dQ / dT or DeltaE / DeltaT, units J/K.
- Specific heat: c = dQ / (m dT), units J/(kg K).
- One-dimensional conduction: Q_energy = k A t DeltaT / Delta x. Heat flux q = k DeltaT / Delta x. 1 W = 1 J/s.
- Thermal conductivity ranking: k_metals > k_ceramics >> k_polymers. Glasses are lower than crystalline ceramics because phonons propagate poorly without a crystalline lattice. Porosity lowers k because air is a poor conductor.
- Diamond has very high k because it has strong bonds, high E, low defects, similar atomic masses, and a simple crystalline cell.
- Free thermal expansion: DeltaL = alpha DeltaT L0 and thermal strain e_th = alpha DeltaT.
- Constrained thermal stress: sigma = E alpha DeltaT in magnitude. Heating a constrained rod produces compression; cooling produces tension.
- Typical CTE order: alpha_ceramics < alpha_metals << alpha_polymers. Stronger bonds give lower alpha, higher E, and higher melting temperature.
- Thermal shock resistance first approximation: TSR about sigma_f k / (E alpha). Best resistance needs high strength, high k, low E, and low alpha; in practice, the main design target is high k and low alpha.
""")
        + section("Phase Diagram Method", """
- Component: element or compound needed to define the system. Binary system: two components.
- Phase: homogeneous portion with uniform chemical and physical characteristics.
- Phase diagram axes for binary condensed systems: temperature on y-axis, composition wt.% on x-axis.
- Gibbs rule: V = C - f + n. For binary condensed diagrams at constant pressure, V = C - f + 1 = 3 - f.
- Horizontal rule/tie line: in a two-phase field, draw a horizontal line at the temperature. Intersections with the phase boundaries give the compositions of the phases.
- Lever rule: the fraction of a phase is proportional to the opposite segment of the tie line. For phases at C1 and C2 and alloy composition C: fraction phase at C1 = (C2 - C) / (C2 - C1); fraction phase at C2 = (C - C1) / (C2 - C1).
- Always state: phases present, phase compositions, relative amounts, and microstructure if requested.
- Eutectic reaction: liquid transforms into two solid phases, L <-> alpha + beta. Peritectic reaction: liquid plus one solid transforms into another solid, L + delta <-> gamma in the Fe-C context.
""")
        + section("Fe-C Diagram Recipes", """
![Fe-C reference diagram](../assets/exercise8_fec_page-03.png)

Key Fe-Fe3C landmarks used in the tutorials:
- Eutectoid composition: about 0.76 wt.% C.
- Eutectoid temperature: about 727 C.
- Maximum C solubility in ferrite at eutectoid: about 0.022 wt.% C.
- Cementite composition: 6.70 wt.% C.
- Eutectic composition: about 4.30 wt.% C at 1147 C.
- Maximum C in austenite at eutectic temperature: about 2.14 wt.% C.

Hypo-eutectoid steel, C0 < 0.76:
- Just above 727 C: alpha ferrite + gamma austenite.
- Pro-eutectoid ferrite fraction = (0.76 - C0) / (0.76 - 0.022).
- Austenite fraction just before eutectoid = (C0 - 0.022) / (0.76 - 0.022).
- At room temperature, that austenite becomes pearlite, so pearlite fraction equals the austenite fraction just before the eutectoid transformation.
- Example from Exercise 8: for 0.20% C, pro-eutectoid ferrite = (0.76 - 0.20)/(0.76 - 0.022) = 76%; pearlite = (0.20 - 0.022)/(0.76 - 0.022) = 24%.

Hyper-eutectoid steel, 0.76 < C0 < 2.14:
- Just above 727 C: gamma austenite + Fe3C cementite.
- Pro-eutectoid cementite fraction = (C0 - 0.76) / (6.70 - 0.76).
- Austenite fraction just before eutectoid = (6.70 - C0) / (6.70 - 0.76).
- At room temperature, that austenite becomes pearlite.
- Example from Exercise 8: for 1.20% C, pearlite = (6.70 - 1.20)/(6.70 - 0.76) = 92.5%; pro-eutectoid cementite = (1.20 - 0.76)/(6.70 - 0.76) = 7.4%.

Peritectic practice from Exercise 9:
- For low-carbon compositions near the peritectic, read phase compositions with a horizontal tie line and then apply the lever rule exactly as in generic diagrams.
- Example for 0.34% C at one two-phase point: if delta ferrite is 0.05% C and liquid is 0.40% C, delta fraction = (0.40 - 0.34)/(0.40 - 0.05) = 17%, liquid = 83%.
- Example for 0.13% C at the same tie line: delta fraction = (0.40 - 0.13)/(0.40 - 0.05) = 77%, liquid = 23%.
""")
        + section("Open Answer Templates", """
Heat conduction in solids:
Define k and units. Explain phonons and free electrons. Say electrons dominate in metals, phonons dominate in ceramics, amorphous materials and polymers conduct poorly, porosity lowers k, and diamond is high-k because of strong bonds, low defects, similar atomic mass, and simple crystal structure.

Sintering:
Define it as densification of compacted powders during firing below the melting point of the main material. Explain diffusion, neck formation between particles, pore reduction, shrinkage, higher strength, and why it is central for ceramics that are not produced by melting.

Safety glass:
Start from brittle glass and defects. Then explain the design loads, tempered glass surface compression/core tension, chemical tempering by ion exchange, laminated/multilayered glass with polymer sheet retaining fragments, armoured glass with metallic network for fire/flame delay, and foamed glass for insulation rather than impact resistance.

Thermoplastic vs thermoset:
Thermoplastics soften on heating and can be reshaped/recycled; chains are linear or branched with secondary bonding. Thermosets cure into a cross-linked network with covalent cross-links, are permanently set, more rigid/brittle/dimensionally stable, and cannot be remelted. Elastomers have few bridges and very high elastic deformation.
"""),
    ),
})


GUIDES.append({
    "slug": "00_introduction",
    "title": "Lecture 00 - Introduction and Exam Strategy",
    "source": "STM/Material(1)/00_Introduction STM 2025-2026.pdf",
    "md": guide(
        "Lecture 00 - Introduction and Exam Strategy",
        "STM/Material(1)/00_Introduction STM 2025-2026.pdf",
        section("What This Lecture Really Tells You", """
- STM gives the non-metallic and general materials foundation: metals, polymers, ceramics, glasses, composites, mechanical properties, and thermal properties.
- The integrated STM-TMM course has one final grade, so phase diagrams and metallic materials still matter, but this STM pack focuses on the Baino/Ferraris/De la Pierre/Martorino materials.
- The final exam simulation shows the professor expects both calculations and explanatory theory, not only memorized definitions.
""")
        + section("How To Use This Open-Note Pack", """
- For a calculation, go first to the master formula sheet and then to the relevant lecture guide.
- For a theory/open question, answer in the sequence: definition, mechanism, property consequence, example/application, and trap/limitation.
- For multiple choice, check whether the statement confuses material classes: e.g. metals conduct by electrons, ceramics/glasses mainly by phonons, polymers by molecular vibrations/rotations; glass is amorphous but still belongs formally to ceramic materials.
""")
        + section("High-Risk Map", """
- Tensile graph interpretation: Lectures 02-03 and Exercises 1-2.
- Creep/fatigue/fracture/hardness: Lecture 03 and Exercise 3.
- Thermal conduction, expansion, stress, thermal shock: Lecture 04 and Exercise 4.
- Generic phase diagrams and lever rule: Lecture 05 and Exercises 5-7.
- Fe-C diagram: Lecture 06 and Exercises 8-9.
- Sintering and ceramic porosity: Lecture 07 and exam simulation.
- Safety glass: Lecture 08.
- Thermoplastic/thermoset comparison: Lecture 09 and exam simulation.
- Rule of mixtures and reinforcement mechanisms: Lecture 10.
"""),
        "00_Introduction_STM_2025_2026.txt",
        12,
    ),
})


GUIDES.append({
    "slug": "01_classification_crystal_defects",
    "title": "Lecture 01 - Classification, Crystal Structures, Defects",
    "source": "STM/Material(1)/01_Classification_crystal_defects 2025-2026.pdf",
    "md": guide(
        "Lecture 01 - Classification, Crystal Structures, Defects",
        "STM/Material(1)/01_Classification_crystal_defects 2025-2026.pdf",
        section("Core Idea", """
Materials science links structure to properties. The professor's chain is: chemical bond and composition -> crystal/amorphous structure -> defects and microstructure -> mechanical, thermal, optical, electrical, technological performance. If asked to choose or explain a material, always connect class, bond, structure, and property.
""")
        + section("Material Classes and Bonding", """
- Metals and alloys: metallic bonding with delocalized electrons. Consequences: good electrical and thermal conductivity, opacity/shine, high plasticity, high toughness, corrosion/oxidation issues, and preparation commonly by melting. Examples include Fe-C steel, Cu-Zn brass, Sn-Pb soldering alloys, Al-based light alloys.
- Ceramics: compounds such as oxides, carbides, nitrides. Ionic/covalent or mixed bonds are strong and often directional. Consequences: high hardness, high melting point, good insulation, brittleness, difficult machining, usually produced by sintering rather than melting.
- Glasses: oxide-based amorphous materials prepared by melting and quenching. They are hard and brittle, insulating, transparent, and soften progressively instead of melting sharply.
- Polymers: organic macromolecules with covalent bonds along chains and weak secondary bonds between chains. Consequences: low density, easy processing, insulation, flexibility, low thermal/mechanical resistance.
- Composites: matrix plus second phase/reinforcement. Classify by matrix (metal, polymer, ceramic, glass/glass-ceramic) or by reinforcement morphology (fibers, particles).
""")
        + section("Crystalline, Amorphous, Semi-Crystalline", """
- Crystalline materials have atoms/ions/molecules in ordered positions in cells. Unit-cell parameters a, b, c and angles alpha, beta, gamma identify the structure.
- Seven basic lattice systems plus centered variants form Bravais lattices. The lecture highlights FCC, BCC, and hexagonal compact structures. Examples: FCC Cu/Al/Ag/Au; BCC Fe/W/Cr; HC Mg/Ti/Zn.
- Amorphous materials have only short-range order. Glasses and many polymers have no exact melting temperature; they soften progressively and are sometimes described as undercooled liquids.
- Semi-crystalline materials include many polymers and glass-ceramics.
- Allotropy/polymorphism means a crystalline solid changes crystal structure while still solid, with property changes in volume, thermal expansion, conductivity, E, density, etc. Fe, Ti, and SiO2 are examples.
""")
        + section("X-Ray and Miller Index Formula Box", """
- Bragg law: n lambda = 2 d sin theta. n is an integer, lambda is wavelength, d is interplanar spacing, theta is the incident angle.
- For cubic cells: d_hkl = a / sqrt(h2 + k2 + l2).
- Miller indexes (h k l) identify crystal planes. XRD distinguishes crystalline from amorphous materials; broad peaks indicate smaller crystal size and amorphous material gives broad diffuse response.
""")
        + section("Point Defects", """
- Point defects are atomic-size defects but perturb a region around them. They create local stress because atoms are not in their correct lattice positions.
- Vacancies: empty lattice sites. They allow atomic motion from site to site, so they enable diffusion. Vacancy concentration increases with temperature by Arrhenius behavior: Nv/Ntot = C exp(-Ev/kT). At 0 K it would be zero, but absolute zero is unreachable.
- Interstitial atoms: atoms crowded into interstitial positions, not lattice knots. Metallic lattices commonly accept H, C, N, O as interstitial impurities. Interstitial solid solutions have a solubility limit; beyond it a second phase forms.
- Steel is the key example: C in Fe is an interstitial solid solution up to the solubility limit and improves mechanical properties by lattice distortion/internal stress.
- Substitutional atoms occupy lattice knots and can be smaller or larger than host atoms, creating tensile or compressive local stress.
- Hume-Rothery rules for substitutional solid solutions: similar atomic radius (difference below about 15%), same crystal structure, similar electronegativity, and compatible valence.
- Ionic solids can have Frenkel defects (vacancy plus interstitial ion) and Schottky defects (paired cation and anion vacancies), increasing electrical conductivity.
""")
        + section("Dislocations and Plastic Deformation", """
- Dislocations are line defects caused by lattice distortion. They are produced during solidification and plastic deformation and are normally present in commercial materials.
- Edge dislocation: extra half-plane of atoms. Screw dislocation: spiral distortion around a dislocation line. Real dislocations are often mixed.
- Plastic deformation in metals is possible because dislocations move. Perfect-crystal sliding would require breaking many bonds at once, giving a theoretical resistance far above real crystals.
- Elastic deformation stretches bonds but does not move dislocations. Plastic deformation is irreversible because applied stress moves dislocations.
- Work hardening: during plastic deformation, dislocations generate and accumulate. Their stress fields interact, block further motion, raise strength/yield stress, and reduce ductility.
- Dislocations also interact with point defects, which is why alloys are stronger than pure metals.
- Covalent solids and ionic solids are brittle because bond directionality or charge repulsion makes dislocation motion difficult. Metallic bonding is non-directional, so dislocations move smoothly and metals are ductile.
""")
        + section("Grain Boundaries and 3-D Defects", """
- Monocrystals have all cells oriented the same way; polycrystals consist of grains, each with its own orientation, separated by grain boundaries.
- Grain boundaries are surface defects, roughly a few atomic diameters thick, where crystalline order is disrupted.
- Fine grain size generally increases mechanical strength. Grain boundaries also have high energy and can be more chemically reactive than grains.
- 3-D defects include precipitates/second phases when solubility limits are exceeded. Pores, cavities, and cracks are macroscopic defects that strongly reduce mechanical, electrical, and optical properties.
- Strengthening methods in this lecture: substitutional solid solution, interstitial solid solution, grain size reduction, second phase/precipitates. The shared mechanism is lattice distortion and obstacles to dislocation motion.
""")
        + section("Amorphous Materials and Viscosity", """
- Amorphous materials have local short-range order but no long-range periodicity. They are metastable and can crystallize if properly heated.
- Glass viscosity decreases continuously with temperature, unlike crystalline materials with a sharp melting point.
- Viscosity unit: 10 Poise = 1 Pa s. Example from the slides: water at room temperature has eta about 10^-3 Pa s, while silica at 1720 C has eta about 10^6 Pa s.
- Transparency trap: glass and diamond can be transparent for different structural reasons; glass is amorphous with no long-range grain-boundary scattering, while diamond is crystalline but can be transparent because visible light does not have the right interaction with its bonding electrons.
""")
        + section("Closed-Question Traps", """
- E and melting temperature follow bond strength; dislocation density controls plastic deformation and strength, not elastic modulus.
- If an interstitial atom is added above solubility, the result is not "more interstitial solution only"; a second phase forms.
- Crystalline does not mean defect-free. A perfect crystal is impossible at finite temperature because vacancies exist.
- Polymers are not purely covalent solids in behavior: covalent bonds are along chains, weak secondary bonds are between chains.
- Glass is amorphous and lacks a defined melting point, but it is formally treated as a ceramic sub-family in this course.
"""),
        "01_Classification_crystal_defects_2025_2026.txt",
        10,
    ),
})


GUIDES.append({
    "slug": "02_mechanical_properties_part_1",
    "title": "Lecture 02 - Mechanical Properties Part 1",
    "source": "STM/Material(1)/02_Mechanical prop part 1 2025-2026.pdf; STM/Material(2)/Exercise1_text.pdf; Exercise2_solved_Ferraris.pdf",
    "md": guide(
        "Lecture 02 - Mechanical Properties Part 1",
        "STM/Material(1)/02_Mechanical prop part 1 2025-2026.pdf; STM/Material(2)/Exercise1_text.pdf; Exercise2_solved_Ferraris.pdf",
        section("Core Idea", """
Mechanical properties are used for material selection and design. The first mechanical lecture builds the elastic/plastic language: stress, strain, Young modulus, Poisson ratio, yield strength, and how to read a tensile curve.
""")
        + section("Stress and Strain", """
- Normal stress: sigma = F/A. In tensile/compression tests use initial area A0 unless the question asks true stress.
- Engineering strain: e = (l - l0)/l0 = Delta l/l0. A strain of 0.002 is 0.2%.
- Shear stress: tau = F/A0. Shear strain: gamma = tan theta.
- The lecture also labels bending and torsion as stress states. For this exam, most calculations use axial stress/strain and simple sections.
""")
        + section("Elastic Behavior and Young Modulus", """
- Hooke law: sigma = E e. E is Young modulus/elastic modulus, units MPa or GPa.
- Elastic deformation is recovered immediately after unloading because bonds return to equilibrium length.
- E is proportional to bond energy and melting temperature. Ranking: E_ceramics > E_metals >> E_polymers.
- E is not related to dislocation motion. Low alloying in steels barely changes E; ordinary steel stays near 210 GPa.
- Amorphous materials generally show lower E than comparable crystalline materials.
- Composite E can be estimated by a volumetric mixture rule: Ec = E1 fv1 + E2 fv2 + ... when the loading geometry supports that approximation.
""")
        + section("Poisson Ratio and Elastic Design", """
- Poisson coefficient: nu = - transverse strain / longitudinal strain. Transverse and longitudinal strains have opposite signs in uniaxial tension/compression.
- Typical values: metals about 0.33; ceramics 0.17-0.27; polymers 0.33-0.5; always below 0.5.
- If a cylinder in tension must reduce diameter by Delta d, compute transverse strain Delta d/d0, then longitudinal strain = - transverse strain / nu, then sigma = E e, then F = sigma A.
""")
        + section("Plastic Deformation and Yield", """
- When stress is high enough and strain exceeds the proportional elastic region, Hooke law no longer applies. Some bonds break and reform in different positions; after unloading, only the elastic part is recovered.
- Yield strength is the frontier between elastic and plastic behavior.
- In steels, interstitial C pins dislocations. The upper yield point corresponds to first unpinning; the lower yield point is the lower stress needed after dislocations move more freely.
- Some metals/alloys do not show a sharp yield point; then use the 0.2% offset yield strength.
""")
        + section("Tutorial Method: Elastic Tensile Calculations", """
- Maximum safe load with safety factor: choose sigma_allowed = YS/n, compute area, then Fmax = sigma_allowed A. Exercise 2 example: Al cylinder d = 10 mm, YS = 150 MPa, n = 2 gives A = 78.5 mm2, sigma_allowed = 75 MPa, Fmax = 5890 N.
- Required diameter for allowed elongation: e = Delta l/l0, sigma = E e, A = F/sigma, d = 2 sqrt(A/pi).
- Comparing material elongation under same F, A, and l0: sigma is the same, e = sigma/E, so lower E gives larger elongation. The exercise compares steel, bronze, and polymer: the polymer elongates enormously because E is very low.
- True fracture vs engineering fracture: engineering fracture uses initial area; true fracture uses final area at necking/rupture. If final diameter is one third of the initial diameter, final area is one ninth of initial area, so true fracture stress is much higher.
""")
        + section("Exam Traps", """
- Do not mix MPa and Pa without converting. 1 N/mm2 = 1 MPa.
- A 0.2% offset strain is 0.002, not 0.2.
- If the question asks "without plastic deformation", compare to yield strength, not tensile strength.
- Elastic modulus controls elastic elongation; yield strength controls onset of plastic deformation; tensile strength is the maximum engineering stress.
- Alloying and heat treatment can change yield and tensile strength a lot but only slightly change E.
"""),
        "02_Mechanical_prop_part_1_2025_2026.txt",
        14,
    ),
})


GUIDES.append({
    "slug": "03_mechanical_properties_part_2",
    "title": "Lecture 03 - Mechanical Properties Part 2",
    "source": "STM/Material(1)/03_Mechanical prop part 2 2025-2026.pdf; STM/Material(2)/Exercise3_text.pdf",
    "md": guide(
        "Lecture 03 - Mechanical Properties Part 2",
        "STM/Material(1)/03_Mechanical prop part 2 2025-2026.pdf; STM/Material(2)/Exercise3_text.pdf",
        section("Core Idea", """
This lecture adds test methods and failure modes: true/engineering curve, necking, compression and bending, creep, fatigue, fracture toughness, hardness, toughness, resilience, Charpy, and ductile-to-brittle transition.
""")
        + section("Tensile Curve Details", """
- Engineering tensile curve uses original area and original length. After necking, engineering stress can drop even though local true stress in the neck continues to increase.
- Necking sequence: localized plastic instability, pore formation, pore coalescence, crack propagation, fracture.
- Ductility by elongation: A% = 100(lf - l0)/l0.
- A standard ductile metal shows elastic region, yield, plastic region, tensile strength, necking, and fracture. Brittle materials fracture with little or no necking.
""")
        + section("Compression and Bending Tests", """
- Compression is important for brittle construction materials such as brick, concrete, and ceramics. It can overestimate performance under real complex stresses because flaws are not opened as in tension.
- Compression of ductile samples can show barreling.
- Bending/flexural tests are often used for brittle materials. In bending, the top surface is in compression and the bottom surface is in tension; failure generally starts on the tensile side.
- Three-point and four-point bending allow flexural strength calculation from failure load and specimen dimensions.
""")
        + section("Creep", """
- Creep is plastic deformation under constant load even when the nominal stress is in the elastic field.
- It is thermally activated: metals above about 0.3-0.4 Tm, ceramics above 0.6-0.7 Tm, amorphous materials above Tg.
- Creep test: constant load, constant high temperature, record strain vs time.
- Primary creep: elastic plus plastic deformation; strain rate decreases due to hardening.
- Secondary creep: steady-state strain rate is low and almost constant because work hardening and annealing balance.
- Tertiary creep: voids/microcracks form, cracks propagate, strain rate accelerates, rupture follows.
- Higher stress and higher temperature shorten creep life. Large grains or single crystals improve creep resistance because there is less grain boundary sliding.
""")
        + section("Fatigue", """
- Fatigue is failure under cyclic loading. It can occur below yield strength and below tensile strength, so it is dangerous even for ductile materials.
- S-N/Wohler curve: stress S versus number of cycles N to failure. Use it to read fatigue life at a given stress or fatigue strength at a given life.
- Stress amplitude: sigma_a = (sigma_max - sigma_min)/2.
- Fatigue limit: below this stress, some materials do not fail within the considered cycle range. Fe, Ti, and steels can show a limit; Al and Cu alloys are treated differently and may fail after enough cycles.
- Fatigue fracture surfaces show a smoother fatigue propagation region and a rough final catastrophic failure region.
- Improve fatigue resistance by surface strengthening, coatings/environmental protection, better design to avoid notches, and fatigue/fracture mechanics life prediction.
""")
        + section("Fracture Mechanics", """
- Macro-defects such as pores and cracks are stress concentrators. The local stress at the crack tip is higher than nominal stress.
- Elliptical crack approximation: sigma_m = 2 sigma_0 sqrt(a/r_t). Larger cracks and sharper crack tips are worse.
- Fracture toughness: KIc = Y sigma_c sqrt(pi a). It combines critical stress, crack size, and geometry. Failure occurs when applied K reaches KIc.
- Ductile fracture has plastic deformation at the crack tip and slow/stable crack propagation. Brittle fracture has little plasticity and fast/catastrophic propagation.
- Weibull/statistical scale effect: larger specimens have higher probability of containing a critical flaw, so they can fail at lower stress. Glass fibers can be much stronger than glass slabs because their volume/defect population is smaller.
""")
        + section("Hardness, Toughness, Resilience, Charpy", """
- Hardness is resistance to localized compressive indentation. It can be measured with spherical or pyramidal indenters.
- Brinell: HB = F/(pi D d). Vickers: HV = 1.854 F/L2. Rockwell measures penetration depth; values outside the scale range are not acceptable.
- Hardness is useful because it is easier and less destructive than tensile testing. For steel, TS is roughly 3.45 HB.
- Toughness is total energy absorbed up to fracture, the whole stress-strain area.
- Resilience is elastic energy absorption. Ur = 1/2 sigma_y e_y = sigma_y2/(2E).
- Charpy test measures absorbed impact energy of a notched sample. It is used to study ductile-to-brittle transition.
- DBTT matters especially for BCC metals such as Fe at low temperature. FCC metals such as Cu and Ni are generally more ductile over temperature.
""")
        + section("Tutorial Method: Creep and Fatigue Questions", """
- If asked to draw creep curves for different temperatures or stresses, keep the same three stages but shift higher T or higher stress to larger strain and shorter rupture time.
- If asked about microstructure in creep: monocrystalline or large-grain materials resist creep better than fine-grained polycrystals because grain-boundary sliding is reduced.
- If asked to use fatigue curves: compute stress from load and area first, then read N from the S-N curve. With a safety coefficient, reduce allowable stress or increase required diameter accordingly.
"""),
        "03_Mechanical_prop_part_2_2025_2026.txt",
        12,
    ),
})


GUIDES.append({
    "slug": "04_thermal_properties",
    "title": "Lecture 04 - Thermal Properties",
    "source": "STM/Material(1)/04_thermal properties_2025-2026.pdf; STM/Material(2)/Exercise4_Q_solved_Ferraris.pdf",
    "md": guide(
        "Lecture 04 - Thermal Properties",
        "STM/Material(1)/04_thermal properties_2025-2026.pdf; STM/Material(2)/Exercise4_Q_solved_Ferraris.pdf",
        section("Core Idea", """
Thermal properties describe how a material responds to heat: temperature rise, dimensional change, energy transport, melting/softening, and thermal stresses. The exam simulation specifically asks heat conduction mechanisms, k ranking, diamond, and thermal expansion/stress calculations.
""")
        + section("Heat Capacity and Specific Heat", """
- Heat capacity C is the ratio between heat added and temperature increase: C = dQ/dT or DeltaE/DeltaT, units J/K.
- Specific heat c is heat capacity per unit mass: c = dQ/(m dT), units J/(kg K). cp is at constant pressure, cv at constant volume; at room temperature the difference is small for many solids.
- Atomic/molecular heat is per mole and is often about 25 J/(mol K) for many solid elements at sufficiently high temperature.
- During phase transformations, heat can be spent on the transformation without increasing temperature.
""")
        + section("Thermal Conductivity", """
- Thermal conductivity is the ability to transfer thermal energy from a hotter region to a colder region.
- Formula for 1-D conduction: Q_energy = k A t DeltaT / Delta x. Heat flux q = k DeltaT / Delta x.
- k has units W/(m K). A is area, t is time, Delta x is thickness, DeltaT is temperature difference.
- Mechanisms: free electrons and phonons. k_total = k_lattice + k_electrons.
- Metals: free-electron contribution dominates, so k is high and related to electrical conductivity.
- Ceramics: phonon contribution dominates; usually lower k than metals.
- Glasses: lower k than crystalline ceramics because vibrational waves propagate poorly without long-range order.
- Polymers: very low k because heat is dissipated by vibration/rotation of molecular chains; amorphous polymers are especially low.
- Porous materials have low k because air in pores is a poor heat conductor. Foams and aerogels exploit this.
""")
        + section("Useful k Values and Ranking", """
- Exercise values: Cu 398 W/mK, Al 247, Fe 81, steel 52, stainless steel 16.
- Ceramics/glasses examples: MgO 38, Al2O3 30, MgAl2O4 15, SiO2 glass about 2, soda-lime glass about 1.7.
- Polymers examples: polyethylene 0.4, nylon 0.24, PTFE 0.25.
- Ranking to state in theory answers: k_metals > k_ceramics >> k_polymers, with glasses below crystalline ceramics.
- Diamond has very high k because of strong bonds, high E, low defects, similar atomic masses, and simple crystalline cell.
""")
        + section("Thermal Expansion", """
- On heating, atoms vibrate more and average bond distance increases, so solids expand.
- Thermal strain: e_th = DeltaL/L0 = alpha DeltaT.
- Linear expansion: DeltaL = alpha DeltaT L0.
- alpha units: K^-1 or C^-1. It varies somewhat with temperature.
- Typical ranges from slides/exercise: polymers about 50-400 x 10^-6 K^-1 (exercise simplified 100 x 10^-6 K^-1), metals about 5-25 or 15-30 x 10^-6 K^-1, glasses about 0.5-20 x 10^-6 K^-1, ceramics about 0.5-15 or 3-7 x 10^-6 K^-1.
- Strong bonds give high E, high melting temperature, and low thermal expansion coefficient.
""")
        + section("Thermal Stress and Thermal Shock", """
- If a body is free, it expands or contracts and stress may remain low.
- If expansion/contraction is constrained, thermal stress develops: |sigma| = E alpha |DeltaT|.
- Heating a constrained rod wants to elongate, so the constraint produces compression. Cooling a constrained rod wants to contract, so the constraint produces tension.
- Joined materials with different CTE develop residual stresses during cooling. If alpha_B > alpha_A, material A is in compression after cooling; if alpha_B < alpha_A, material A is in tension.
- Thermal shock resistance first approximation: TSR about sigma_f k / (E alpha). High k and low alpha are favorable; high strength helps; brittle materials often fracture under thermal stresses.
""")
        + section("Tutorial Methods", """
- Conducted energy: convert thickness to meters, area to m2, time to seconds, then Q = k A t DeltaT / Delta x.
- Example from Exercise 4: polymer component, k = 0.58 W/mK, DeltaT = 30 K, Delta x = 0.002 m, A = 1 mm2 = 1e-6 m2, t = 3600 s. Heat flux = 8700 W/m2, energy = 31.32 J.
- Expansion example: steel alpha = 14 x 10^-6 C^-1, L0 = 50 m, DeltaT = 30 C. DeltaL = 14e-6 x 30 x 50000 mm = 21 mm.
- Constrained stress example: steel E = 207 GPa, alpha = 12e-6 C^-1, DeltaT = 60 C. sigma = 207000 MPa x 12e-6 x 60 = 149 MPa compression.
- Temperature limit example: with constrained brass, sigma = E alpha DeltaT; solve DeltaT = sigma/(E alpha), then add to initial temperature.
"""),
        "04_thermal_properties_2025_2026.txt",
        14,
    ),
})


GUIDES.append({
    "slug": "05_phase_diagrams",
    "title": "Lecture 05 - Phase Diagrams",
    "source": "STM/Material(1)/05_phase diagrams_2025-2026.pdf; STM/Material(2)/Exercises5-7",
    "md": guide(
        "Lecture 05 - Phase Diagrams",
        "STM/Material(1)/05_phase diagrams_2025-2026.pdf; STM/Material(2)/Exercises5-7",
        section("Core Idea", """
Phase diagrams are equilibrium maps. They show which phases exist at a given composition and temperature, what the phase compositions are, and what relative amounts are present. The professor expects you to read the diagram, not memorize an isolated formula.
""")
        + section("Definitions", """
- Component: element or chemical compound whose presence is necessary and sufficient to define the system.
- Alloy: mixture of two or more elements with metallic properties.
- Binary system: two components.
- Phase: homogeneous part of a system with uniform chemical and physical characteristics.
- Monophasic system: one phase. Biphasic system: two phases differing by state, composition, and/or crystal structure.
- Microstructure and properties depend on composition, temperature, and pressure; in binary condensed diagrams pressure is normally treated as constant.
""")
        + section("How To Read Any Binary Diagram", """
1. Locate the alloy composition on the x-axis and the temperature on the y-axis.
2. Identify the field: one phase or two phases.
3. In a one-phase field, the phase composition equals the alloy composition and the relative amount is 100%.
4. In a two-phase field, draw a horizontal tie line. The intersections with the boundaries give the compositions of the two phases.
5. Apply the lever rule for relative amounts.
6. If requested, describe the microstructure: primary phase, eutectic mixture, pro-eutectoid phase, precipitates, etc.
""")
        + section("Gibbs Rule", """
- General form: V = C - f + n.
- For condensed binary diagrams at constant pressure: V = C - f + 1 = 3 - f.
- One-phase field in binary diagram: f = 1, V = 2. Temperature and composition can vary.
- Two-phase field: f = 2, V = 1. If temperature is fixed, phase compositions are fixed by the tie line.
- Three phases in binary condensed system: f = 3, V = 0. This is invariant, such as eutectic/peritectic reactions at fixed T and compositions.
""")
        + section("Horizontal Rule and Lever Rule", """
- Horizontal/tie-line rule: draw a horizontal line across the two-phase region at the temperature of interest. The left and right boundary compositions are the phase compositions.
- Lever rule formula: for phases at C1 and C2 and alloy at C, fraction of phase at C1 = (C2 - C)/(C2 - C1), fraction of phase at C2 = (C - C1)/(C2 - C1).
- Memory trick: amount of a phase is the opposite length of the tie line.
- Always multiply by 100 for percentage.
""")
        + section("Diagram Families", """
- Complete solubility in liquid and solid state: examples Cu-Ni or Ag-Au. There is a liquidus, solidus, L, solid solution, and L+solid region.
- Complete solubility in liquid but no solubility in solid: eutectic diagram with L transforming into A+B or alpha+beta.
- Partial solubility in solid: eutectic diagram with alpha, beta, alpha+beta, L+alpha, L+beta fields and solvus lines.
- Peritectic: L + delta <-> gamma.
- Congruent melting compound: compound melts to liquid of same composition.
- Non-congruent melting compound: compound decomposes/transforms with liquid plus another phase.
""")
        + section("Microstructure Language", """
- Primary phase: phase that forms before an invariant reaction.
- Eutectic structure: fine mixture formed by L -> alpha + beta at eutectic temperature/composition.
- Hypoeutectic vs hypereutectic: composition respectively below or above eutectic composition. Describe primary phase plus eutectic constituent.
- During cooling below a solvus, phase compositions change along solvus curves and relative amounts change by lever rule.
""")
        + section("Tutorial Checklist", """
- Exercises 5-7 repeatedly ask: determine phases, compositions, relative amounts, and draw/schematize microstructure.
- For any point labeled at the beginning/end of a eutectic transformation, be precise: just before the reaction, liquid may still be present; just after it, the liquid has transformed into the eutectic mixture.
- If the question asks "complete solubility" vs "partial/no solubility", do not invent extra phases. Use the region labels from the diagram family.
""")
        + section("Exam Traps", """
- Do not use the alloy composition as phase composition in a two-phase field. Use the tie-line endpoints.
- Do not apply lever rule in a single-phase field.
- Do not say a phase diagram is a kinetic cooling curve. These diagrams are equilibrium diagrams and assume slow cooling.
- Do not forget that phase diagrams have temperature on y-axis and composition wt.% on x-axis in this course.
"""),
        "05_phase_diagrams_2025_2026.txt",
        12,
    ),
})


GUIDES.append({
    "slug": "06_phase_diagram_examples_fe_c",
    "title": "Lecture 06 - Phase Diagram Examples and Fe-C",
    "source": "STM/Material(1)/06_phase diagrams examples_2025-2026.pdf; STM/Material(2)/Exercises8-9",
    "md": guide(
        "Lecture 06 - Phase Diagram Examples and Fe-C",
        "STM/Material(1)/06_phase diagrams examples_2025-2026.pdf; STM/Material(2)/Exercises8-9",
        section("Fe-C Reference Diagram", """
![Fe-C reference diagram](../assets/exercise8_fec_page-03.png)

Keep these landmarks in your notes:
- 0.022 wt.% C: approximate maximum solubility of C in ferrite at the eutectoid temperature.
- 0.76 wt.% C: eutectoid composition.
- 727 C: eutectoid temperature.
- 2.14 wt.% C: approximate maximum solubility of C in austenite at the eutectic temperature.
- 4.30 wt.% C: eutectic composition.
- 6.70 wt.% C: cementite Fe3C composition.
""")
        + section("Steel Microstructures at Room Temperature", """
- Hypo-eutectoid steel, C < 0.8% in the slides: ferrite and pearlite.
- Eutectoid steel, C about 0.8%: pearlite.
- Hyper-eutectoid steel, C > 0.8%: cementite and pearlite.
- Pearlite is the eutectoid mixture of ferrite + cementite formed from austenite.
- Coarse pearlite forms at higher isothermal transformation temperature because diffusion is easier. Fine pearlite forms at lower temperature because diffusion distance is shorter and lamellae are finer.
- Spheroidite/globular pearlite forms by heating pearlitic steel just below eutectoid temperature for a long time, for example 700 C for 18-24 h. Driving force: reduce ferrite/cementite interfacial area. Result: easier plastic deformation.
""")
        + section("Fe-C Lever Rule Recipes", """
Hypo-eutectoid, C0 < 0.76:
- Pro-eutectoid ferrite = (0.76 - C0)/(0.76 - 0.022).
- Austenite just before eutectoid = (C0 - 0.022)/(0.76 - 0.022).
- At room temperature, pearlite fraction equals the austenite fraction just before eutectoid.

Hyper-eutectoid, C0 > 0.76:
- Pro-eutectoid cementite = (C0 - 0.76)/(6.70 - 0.76).
- Austenite just before eutectoid = (6.70 - C0)/(6.70 - 0.76).
- At room temperature, pearlite fraction equals the austenite fraction just before eutectoid.

Worked tutorial examples:
- 0.20% C steel: pro-eutectoid ferrite = 76%, pearlite = 24%.
- 1.20% C steel: pearlite = 92.5%, pro-eutectoid cementite = 7.4%.
""")
        + section("Peritectic Exercise Pattern", """
- Peritectic reaction in Fe-C: L + delta -> gamma.
- For low-carbon steels near the peritectic, identify whether the point is L, L+delta, L+gamma, delta+gamma, or gamma.
- Use the horizontal rule for phase compositions, then lever rule for relative amounts.
- Example for 0.34% C at a two-phase L+delta point: if C_delta = 0.05 and C_L = 0.40, delta = (0.40 - 0.34)/(0.40 - 0.05) = 17%, liquid = 83%.
- Example for 0.13% C on the same tie line: delta = (0.40 - 0.13)/(0.40 - 0.05) = 77%, liquid = 23%.
""")
        + section("Martensite, Tempering, Bainite", """
- Martensite is produced when austenite transforms without enough diffusion for equilibrium ferrite/cementite formation. It is a hard, brittle, supersaturated phase and is not read as an equilibrium field on the Fe-C phase diagram.
- Tempering martensite decomposes martensite into ferrite plus fine cementite particles. Tempered martensite improves toughness relative to as-quenched martensite.
- Bainite is another non-equilibrium transformation product, typically represented by time-temperature transformation thinking rather than an equilibrium Fe-C diagram.
- Exam trap: phase diagrams show equilibrium phases. Martensite is not located as a stable phase field on the equilibrium Fe-Fe3C diagram.
""")
        + section("Al-Cu/Precipitation Example", """
- The lecture examples include precipitation hardening. If solubility decreases on cooling, a supersaturated solid solution can precipitate fine particles.
- More solute can give a higher density of precipitates, increasing strength when precipitates obstruct dislocation motion.
- Overaging/coarsening lowers strengthening because precipitates become fewer/larger and easier for dislocations to bypass.
""")
        + section("Exam Traps", """
- Pearlite is not a phase; it is a microconstituent made of ferrite + cementite.
- Pro-eutectoid ferrite/cementite forms before the eutectoid reaction; pearlite forms from the remaining austenite during the eutectoid reaction.
- Do not use 6.70 as if it were pure carbon; it is cementite Fe3C composition in wt.% C.
- If asked at 400 C for pro-eutectoid cementite in a hyper-eutectoid steel, the amount of pro-eutectoid cementite is still calculated from the 727 C eutectoid tie line; below 727 C the remaining austenite has transformed to pearlite.
"""),
        "06_phase_diagrams_examples_2025_2026.txt",
        18,
    ),
})


GUIDES.append({
    "slug": "07_ceramics",
    "title": "Lecture 07 - Ceramic Materials",
    "source": "STM/Material(1)/07_ceramics 2025-2026.pdf",
    "md": guide(
        "Lecture 07 - Ceramic Materials",
        "STM/Material(1)/07_ceramics 2025-2026.pdf",
        section("Core Idea", """
Ceramics are mostly compounds of metallic and non-metallic elements with ionic, covalent, or mixed bonding. Strong directional bonds explain their high hardness, high refractoriness, chemical inertia, insulation, and brittleness.
""")
        + section("Properties", """
- Good compressive strength.
- Very poor tensile and bending strength.
- Fragility and low toughness.
- High hardness.
- Good heat resistance/refractoriness.
- Good thermal and electrical insulation.
- Good chemical inertia against acid and alkaline attack.
- Poor thermal shock resistance.
- Typical examples in the wider course: bricks, concrete, tiles, porcelains, alumina, zirconia, silicon carbide, silicon nitride.
""")
        + section("Why Ceramics Are Brittle", """
- Ionic/covalent bonds are strong and often directional.
- Dislocation motion is difficult. In ionic crystals, sliding can put equal charges next to each other, causing repulsion and fracture. In covalent solids, directional bonds resist sliding.
- Defects and pores act as stress concentrators. Under tension/bending, cracks open; under compression, cracks are less severe, so compressive strength is much higher.
""")
        + section("Production Cycle", """
- Minerals/chemical process -> powders -> crushing, milling, sieving -> molding/compaction to a green body -> sintering -> densified body -> optional surface treatment -> final product.
- Green body means compacted particles before firing/densification.
- Traditional clay-based materials use firing/sintering above about 1000 C depending on composition and desired properties.
""")
        + section("Molding Methods", """
- Powder pressing: low water in paste, about 5-7%.
- Extrusion: intermediate water, about 15-20%.
- Slip casting: high water, about 35-40%.
- More water improves flowability/forming for slip casting but requires removal during drying/firing.
""")
        + section("Sintering Answer Template", """
Sintering is thermal treatment of compacted powders below the melting point of the main material. Atomic diffusion causes particles to bond, necks to grow between particles, porosity to decrease, density to increase, and mechanical strength to improve. It is central for ceramics because they are difficult to melt/process as liquids and their high melting points make powder routes practical.
""")
        + section("Vitrification and Porosity", """
- Sintering aids can create vitrification: a liquid glass phase forms, flows into pores, and accelerates diffusion compared with solid-state sintering.
- Porosity usually lowers tensile strength because pores are crack/flaw origins and reduce load-bearing area.
- Closed, fine pores can improve thermal insulation because trapped air conducts poorly. This is why exam options saying pores have "no effect" on mechanical properties are suspicious.
- The exam simulation's porosity MCQ has a trap: pores are formed during sintering and can improve thermal insulation, but they do not improve tensile strength.
"""),
        "07_ceramics_2025_2026.txt",
        18,
    ),
})


GUIDES.append({
    "slug": "08_glass_and_safety_glass",
    "title": "Lecture 08 - Glass and Related Products",
    "source": "STM/Material(1)/08_glass 2025-2026.pdf",
    "md": guide(
        "Lecture 08 - Glass and Related Products",
        "STM/Material(1)/08_glass 2025-2026.pdf",
        section("Core Idea", """
Glass is amorphous: no long-range atomic order, but short-range order exists. Silica glass keeps local SiO4 tetrahedra but lacks the ordered long-range quartz structure. Glasses do not have a sharp melting point like crystals; they progressively soften around the glass transition temperature Tg.
""")
        + section("Glass Transition and Formation", """
- Crystalline solids have a melting temperature Tm. Glasses have a glass transition temperature Tg above which they progressively soften.
- Tg can be viewed as the temperature where glass changes from rigid/brittle to rubber-like behavior. In many commercial glasses Tg is roughly 2/3 Tm.
- Glass-forming liquids cool below solidification temperature without crystallization because viscosity becomes so high that random liquid structure is frozen in.
- Glass is metastable; devitrification is crystallization of glass, which can occur over long times or be accelerated by thermal treatment below Tm.
- Any substance that can turn into an amorphous state under suitable conditions can be called a glass: ceramics, amorphous metals, organic compounds, polymers, and some native elements.
""")
        + section("Glass vs Ceramics vs Glass-Ceramics", """
- Fully crystalline ceramics: rocks/stones, cements, alumina, zirconia, etc.
- Glasses: completely amorphous, non-crystalline ceramic sub-family.
- Glass-ceramics: crystalline phases embedded in residual glass. They are made by deliberate devitrification of a parent glass to improve strength, hardness, stiffness, or thermal shock resistance.
""")
        + section("Glass Constituents and Composition", """
- Network formers: oxides such as SiO2 and B2O3 that can form the glassy network.
- Network modifiers: oxides such as Na2O and CaO that do not form the network alone but enter it and break bridging oxygens, producing non-bridging oxygens.
- Intermediates: oxides such as Al2O3 and TiO2 that can enter/stabilize the network.
- Soda-lime-silica glass (SiO2-Na2O-CaO) is low-cost, easy to process/form, stable in water, and used for windows, vehicle glass, windscreens, and light bulbs. Al2O3 increases durability.
""")
        + section("Glass Synthesis and Processing", """
- Most common route: mix reagents in crucible, melt above 1000 C, then quench/form by pouring, pressing, blowing, or drawing into fibers.
- Float glass production uses a molten tin bath to make flat glass sheets.
- Colored glasses are obtained by adding chromophore ions as small amounts of metal oxides.
""")
        + section("Mechanical Behavior and Properties", """
- Glasses are brittle and contain surface/macroscopic defects. Stress concentration near pores/cracks explains low tensile/flexural strength compared with compressive strength.
- Typical common glass properties from slides: density around 2.5 g/cm3, Young modulus about 75 GPa, compressive strength about 1000 MPa, tensile/flexural strength about 40 MPa for common glasses and up to 200 MPa for special glasses, CTE about 9 x 10^-6 K^-1, very low thermal conductivity, refractive index about 1.5.
- Fracture toughness relation: KIc = Y sigma_c sqrt(pi a). Strength depends on defect size, stress, and geometry.
""")
        + section("Safety Glass - Open Answer Block", """
Safety glass is a design response to the brittleness and defect sensitivity of glass. A complete answer should mention the loads and the safety mechanism.

Loads glass sheets can face:
- Dynamic loads: wind, pressure of crowd.
- Static loads: own weight, snow, hydrostatic pressure.
- Accidental loads: hail, vibration, torsion, seismic phenomena.
- Soft-body impact: people/animals.
- Hard-body impact: stones or burglary tools.
- Projectile impact: firearms.
- Thermal shock: fire.

Armoured glass:
- Contains a metallic network inside the glass.
- Mechanical strength is only slightly improved, but in fire the network delays flame propagation and helps keep softened/broken glass in position.

Tempered glass:
- Physical thermal tempering or chemical strengthening creates permanent compressive stress on the surface and tensile stress in the core.
- Surface compression improves impact resistance because cracks must overcome compression before opening.
- Chemical tempering can use ion exchange: larger K+ ions replace smaller Na+ ions near the surface, inducing compressive stress.

Multilayered/laminated glass:
- Combines glass sheets with a polymeric interlayer.
- It has good shock resistance, rupture is localized, and fragments remain attached to the polymer sheet instead of scattering.
""")
        + section("Foamed Glass", """
- Foamed/cellular glass is used for thermal and acoustic insulation. It has a cellular, waterproof, gas-proof structure with many small closed, non-interconnected pores.
- Main features: good insulation, low density, chemical stability, adequate compressive strength.
- Air is a poor heat and sound conductor, so closed pores are useful.
- Foam glass can be made from waste glass, broken windscreens, recycled glass, silica sandstone, and coal powder. Heating in molds around 600-800 C sinters glass powders and gas bubbles expand the mass.
""")
        + section("Static Fatigue and Stress Corrosion", """
- Glass strength can decrease over time. A glass can break after a certain time under a static stress lower than expected.
- This is static fatigue: defect population evolves with time.
- In silica-based glasses/ceramics, water molecules attack deformed Si-O-Si bonds at crack tips. Stress plus chemical attack causes subcritical crack growth even when K < KIc.
- Final reaction idea from slide: bridging oxygen removal interrupts the network and extends the crack.
""")
        + section("Exam Traps", """
- Glass is amorphous but has short-range order.
- Tg is not Tm. Glass softens progressively instead of melting sharply.
- Tempering does not remove all tensile stress; it intentionally creates surface compression and core tension.
- Armoured glass is mainly about holding the sheet in place/flame delay, not huge bending-strength improvement.
- Foamed glass is for insulation, not transparent window panes.
"""),
        "08_glass_2025_2026.txt",
        14,
    ),
})


GUIDES.append({
    "slug": "09_polymers",
    "title": "Lecture 09 - Polymeric Materials",
    "source": "STM/Material(1)/09_Polymers 2025-2026.pdf",
    "md": guide(
        "Lecture 09 - Polymeric Materials",
        "STM/Material(1)/09_Polymers 2025-2026.pdf",
        section("Core Idea", """
Polymers are macromolecular materials. Their properties depend on chain chemistry, molecular weight, architecture, crystallinity, cross-linking, additives, and temperature relative to Tg.
""")
        + section("Basic Terminology", """
- Polymer: many units. Monomer: basic building block. Macromolecule: very large molecule/chain. Plastics is used as a synonym for polymers in the lecture.
- Most polymers have a hydrocarbon backbone with possible functional groups. Silicones are based on -Si-O-Si-O- chains and are used as sealants.
- Advantages: low cost, low electrical/thermal conductivity, low density, high strength-to-weight ratio, corrosion resistance, noise reduction, color/aesthetic flexibility, easy manufacturing, many shapes.
- Limitations: low mechanical properties, low thermal resistance, flammability, damage by visible light/UV.
- Additives include fillers, lubricants, pigments, and carbon black in tires.
""")
        + section("Polymerization", """
- Polymerization combines monomers into covalently bonded chains or networks.
- Addition polymers: monomers are alkenes or substituted alkenes with C=C bonds. Example: polyethylene from ethylene. Chain growth often needs an initiating species/catalyst.
- Condensation polymers: functional groups react, usually with loss of a small by-product such as water. Example: PET from terephthalic acid and ethylene glycol. Condensation can be slower and often needs heat.
- Polar functional groups can increase chain-chain attraction, crystallinity, and tensile strength.
""")
        + section("Molecular Weight, Architecture, Crystallinity", """
- Molecular weight distribution and degree of polymerization strongly affect mechanical properties and viscosity.
- Linear polymers: sequential chains. Branched polymers: more resistance to deformation/load than simple linear chains. Cross-linked polymers/thermosets: higher hardness, strength, stiffness, brittleness, and dimensional stability. Networked polymers: highly cross-linked, strong resistance to radiation/UV/X-rays/electron beams and higher mechanical properties.
- Copolymers contain two monomer types; terpolymers contain three. They are used to combine/modulate properties, such as styrene-butadiene tires or ABS helmets.
- Amorphous polymers have disordered chains. Semi-crystalline polymers have crystallites. Crystallinity can be increased by controlling/decreasing cooling rate.
""")
        + section("Thermoplastic vs Thermoset vs Elastomer", """
- Thermoplastics soften on heating, can be shaped, return to original solid state on cooling, and can be reformed repeatedly. They are recyclable. Examples: acrylics, nylons, polyethylenes.
- Thermosets cure into one giant cross-linked molecule with strong covalent cross-links. They are permanently set, can be formed only the first time, and the curing reaction is irreversible. Examples: resins and adhesives.
- The slide says "No Tg value" for thermosets; for exam wording, follow the professor's statement: thermosets do not behave like reshapeable thermoplastics around Tg because the network is permanently cross-linked.
- Elastomers are linear polymers with a low number of bridging bonds, introduced after molding, giving a three-dimensional structure and very high elastic deformation up to about 800%.
""")
        + section("Temperature and Stress-Strain Behavior", """
- A: thermosets or thermoplastics below Tg behave rigid/brittle.
- B: thermoplastics above Tg become softer/more ductile.
- C: elastomers show very large elastic deformation.
- Semi-crystalline thermoplastics above Tg can show upper and lower yield points and neck propagation. Chains orient in the neck, locally strengthening the material, and elongation proceeds by neck propagation along the gauge length.
""")
        + section("Exam Simulation Answer: Thermoplastic vs Thermoset", """
Thermoplastics consist of chains held together mainly by secondary bonds between chains, so heating can allow chain mobility and reshaping. They can be recycled and remolded. Thermosets have covalent cross-links forming a permanent network after curing; heating does not allow remelting/remolding, so they are more dimensionally stable and often harder/stiffer/more brittle. In a tensile-curve question, thermosets correspond to a more rigid/brittle curve with limited plastic deformation, while thermoplastics above Tg can show yielding and large deformation.
""")
        + section("Traps", """
- Do not say all polymers are fully amorphous; many are semi-crystalline.
- Do not confuse elastomer with thermoplastic: elastomers have few bridges and very large elastic strain.
- Low density does not mean high stiffness; E of polymers is much lower than metals/ceramics.
- Additives can strongly change polymer behavior, so "polymers always..." options are often traps.
"""),
        "09_Polymers_2025_2026.txt",
        14,
    ),
})


GUIDES.append({
    "slug": "10_composites",
    "title": "Lecture 10 - Composite Materials",
    "source": "STM/Material(1)/10_composites 2025-2026.pdf",
    "md": guide(
        "Lecture 10 - Composite Materials",
        "STM/Material(1)/10_composites 2025-2026.pdf",
        section("Core Idea", """
A composite is a combination of two or more chemically distinct and insoluble phases with a recognizable interface, arranged so properties/structural performance exceed the constituents acting independently. Usually one phase is a continuous matrix and another is reinforcement.
""")
        + section("Classification", """
- By matrix: metal matrix composites (MMCs), ceramic matrix composites (CMCs), polymer matrix composites (PMCs/OMCs).
- By reinforcement: fibers (long or short), particles, dispersions.
- Reinforcement can be metal, polymer, ceramic, or glass depending on the composite design.
""")
        + section("Matrix and Bonding Logic", """
- MMCs: matrix is relatively soft/flexible; reinforcement must have high strength/stiffness; reinforcement-matrix bond must be strong because load transfers from matrix to reinforcement.
- PMCs: matrix is relatively soft/flexible; reinforcement must have high strength/stiffness; strong bond is needed for load transfer.
- CMCs: matrix is hard and brittle; reinforcement should have high tensile strength to arrest crack growth; the reinforcement-matrix bond should be relatively weak enough to allow pull-out and crack energy dissipation rather than immediate brittle failure.
""")
        + section("Fiber-Reinforced Polymer-Matrix Composites", """
- Fibers improve strength, fatigue resistance, Young modulus, and strength-to-weight ratio.
- The fiber provides strength and stiffness.
- The matrix binds fibers, transfers load between fibers, and protects them.
- Common PMC fibers: glass fibers, carbon fibers, aramid fibers.
- Common MMC fibers: boron, carbon, oxide ceramic, non-oxide ceramic fibers.
""")
        + section("Reinforcement Mechanisms", """
- Fiber mechanisms: load transfer, debonding, pull-out, crack bridging/arrest.
- Particle mechanisms: crack deviation, crack blocking, increased crack path, particle breaking, and energy consumption.
- The useful mechanism depends on interface strength. A very strong interface is good for stiffness/load transfer in MMC/PMC; a controlled weaker interface can improve toughness in CMCs.
""")
        + section("Rule of Mixtures", """
- The rule of mixtures predicts properties such as density, elastic modulus, electrical conductivity, and thermal conductivity from constituent properties and volume fractions.
- Density: rho_c = rho_f f_f + rho_m f_m, with f_m = 1 - f_f.
- Longitudinal modulus for continuous unidirectional fibers parallel to load: E_c = E_f f_f + E_m f_m.
- Transverse modulus approximation: 1/E_c = f_f/E_f + f_m/E_m.
- Trap: the high parallel modulus only applies when continuous fibers are aligned with the load.
""")
        + section("Processing", """
- Pultrusion: resin-impregnated glass strands are pulled through a die, similar in idea to extrusion, producing pipes, channels, I-beams, etc.
- Resin transfer molding: reinforcement preform is placed in a closed mold; thermoset matrix is injected; after curing, the part is ejected.
""")
        + section("Exam Traps", """
- A composite is not just a mixture: phases must be distinct/insoluble with an interface and improved combined performance.
- Matrix and reinforcement roles are different. Matrix is not necessarily the strongest phase.
- In CMCs, a relatively weak interface can be beneficial for pull-out and crack arrest; do not apply the PMC/MMC strong-interface rule blindly.
- Rule of mixtures must use volume fractions, not weight fractions, unless converted.
"""),
        "10_composites_2025_2026.txt",
        18,
    ),
})


def add_md_lecture_07_if_missing() -> None:
    # Placeholder kept only to make the guide list visually grouped in the source file.
    return None


def iter_markdown_lines(md: str):
    for line in md.splitlines():
        yield line.rstrip()


def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.25


def add_hyperlink_style(doc: Document) -> None:
    # Keep URLs/source paths readable without adding Word hyperlink relationships.
    styles = doc.styles
    if "SourceNote" not in styles:
        st = styles.add_style("SourceNote", 1)
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(9)
        st.font.color.rgb = RGBColor.from_string("555555")
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.15


def add_numbering_compat(paragraph):
    # python-docx's built-in list styles render correctly with Word/LibreOffice.
    # This small tab stop keeps wrapped list text aligned more cleanly.
    pPr = paragraph._p.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        pPr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "left")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)


def md_to_docx(md: str, out_path: Path) -> None:
    doc = Document()
    style_document(doc)
    add_hyperlink_style(doc)

    first_title = True
    for line in iter_markdown_lines(md):
        if not line:
            continue
        if line.startswith("!["):
            m = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if m:
                alt, rel = m.groups()
                img_path = (MD_OUT / rel).resolve() if not Path(rel).is_absolute() else Path(rel)
                if img_path.exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(str(img_path), width=Inches(6.2))
                    cap = doc.add_paragraph(alt)
                    cap.style = "SourceNote"
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue
        if line.startswith("# "):
            if first_title:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(line[2:])
                set_run_font(run, size=22, bold=True, color="0B2545")
                first_title = False
            else:
                doc.add_heading(line[2:], level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_numbering_compat(p)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            p.add_run(line[2:])
            continue
        if line.startswith("Source:"):
            p = doc.add_paragraph(style="SourceNote")
            p.add_run(line)
            continue
        p = doc.add_paragraph()
        p.add_run(line)

    doc.save(out_path)


def main() -> None:
    MD_OUT.mkdir(parents=True, exist_ok=True)
    DOCX_OUT.mkdir(parents=True, exist_ok=True)
    for item in GUIDES:
        md = item["md"]
        md_path = MD_OUT / f"{item['slug']}.md"
        docx_path = DOCX_OUT / f"{item['slug']}.docx"
        md_path.write_text(md, encoding="utf-8")
        md_to_docx(md, docx_path)

    index_lines = [
        "# STM Open-Note Guides Index",
        "",
        "Generated from local STM lecture and exercise PDFs only.",
        "",
        "## Files",
        "",
    ]
    for item in GUIDES:
        index_lines.append(f"- [{item['title']}](markdown/{item['slug']}.md) - DOCX: `docx/{item['slug']}.docx`")
    index_lines.extend([
        "",
        "## Fast Routes",
        "",
        "- Fe-C diagram: `06_phase_diagram_examples_fe_c` and `00_master_formula_and_exam_methods`.",
        "- Safety glass: `08_glass_and_safety_glass`.",
        "- Thermal calculation: `04_thermal_properties` and master formula sheet.",
        "- Tensile/elastic graph calculation: `02_mechanical_properties_part_1` and `03_mechanical_properties_part_2`.",
        "- Creep/fatigue/fracture: `03_mechanical_properties_part_2`.",
        "- Sintering/ceramic porosity: `07_ceramics`.",
        "- Thermoplastic vs thermoset: `09_polymers`.",
        "- Rule of mixtures/composites: `10_composites`.",
        "",
    ])
    (ROOT / "INDEX.md").write_text("\n".join(index_lines), encoding="utf-8")
    print(f"wrote {len(GUIDES)} markdown files and {len(GUIDES)} docx files")


if __name__ == "__main__":
    main()
